import matplotlib.pyplot as plt
from model.Student_Search import *
from docx import Document
from docx.shared import Inches


# draw a pie chart for chinese scores distribution
def pie_district_chinese(students):
    labels = ['scores > 90', '80 < scores < 90', 'scores < 80']
    values = [show_scores_district(students)[0][0], show_scores_district(students)[0][1],
              show_scores_district(students)[0][2]]
    plt.pie(values, labels=labels, autopct="%.2f%%", labeldistance=1.12, textprops=dict(fontsize=16), radius=1.1)
    plt.title('Chinese Scores District', fontsize=20)
    plt.savefig("Chinese_District.png")
    plt.show()


# draw a pie chart for math scores distribution
def pie_district_math(students):
    labels = ['scores > 90', '80 < scores < 90', 'scores < 80']
    values = [show_scores_district(students)[1][0], show_scores_district(students)[1][1],
              show_scores_district(students)[1][2]]
    plt.pie(values, labels=labels, autopct="%.2f%%", labeldistance=1.12, textprops=dict(fontsize=16), radius=1.1)
    plt.title('Math Scores District', fontsize=20)
    plt.savefig("Math_District.png")
    plt.show()


# draw a pie chart for english scores distribution
def pie_district_english(students):
    labels = ['scores > 90', '80 < scores < 90', 'scores < 80']
    values = [show_scores_district(students)[2][0], show_scores_district(students)[2][1],
              show_scores_district(students)[2][2]]
    plt.pie(values, labels=labels, autopct="%.2f%%", labeldistance=1.12, textprops=dict(fontsize=16), radius=1.1)
    plt.title('English Scores District', fontsize=20)
    plt.savefig("English_District.png")
    plt.show()


# generate a report for print use
def report_generate(students):
    pie_district_chinese(students)
    pie_district_math(students)
    pie_district_english(students)

    document = Document()
    document.add_heading('Class Scores Report', 0)

    # student rank
    document.add_heading('Scores Rank', 1)
    scores = [['Name', 'Id', 'Chinese', 'Math', 'English']]
    for student in students:
        scores.append([student.name, str(student.id), str(student.list_scores[0]), str(student.list_scores[1]),
                       str(student.list_scores[2])])
    row = len(scores)
    col = len(scores[0])
    table = document.add_table(rows=row, cols=col)
    table.style = 'Light Shading'
    for r in range(row):
        for c in range(col):
            table.cell(r, c).text = scores[r][c]

    # picture
    document.add_heading('Scores District', 1)
    document.add_picture('Chinese_District.png', width=Inches(3))
    document.add_picture('English_District.png', width=Inches(3))
    document.add_picture('Math_District.png', width=Inches(3))

    # average and median
    document.add_heading('Average and Median', 1)
    info = [
        ['   ', 'Average', 'Median'],
        ['Chinese', str(show_average(students, len(students))[0]), str(show_median(students)[0])],
        ['Math', str(show_average(students, len(students))[1]), str(show_median(students)[1])],
        ['English', str(show_average(students, len(students))[2]), str(show_median(students)[2])]
    ]
    row = len(info)
    col = len(info[0])
    table = document.add_table(rows=row, cols=col)
    table.style = 'Light Shading'
    for r in range(row):
        for c in range(col):
            table.cell(r, c).text = info[r][c]
    # save file
    document.add_page_break()
    document.save('Report_Generated.docx')
